import streamlit as st
from docx import Document
from docx.shared import Inches
from io import BytesIO
import re
from PIL import Image

st.set_page_config(page_title="Smart CV Builder", page_icon="🧠")

st.title("🧠 Smart CV Builder")
st.write("""
Upload your CV Template (Word, Image, or PDF).  
This version supports:
- 📄 `.docx` templates with placeholders like `{{name}}`, `{{summary}}`
- 🖼️ JPEG/PNG/PDF templates
- 👤 User photo upload (`{{photo}}`)
- 🚫 Auto-removal of empty sections
""")

# --- Upload Template (Word, Image, or PDF) ---
uploaded_template = st.file_uploader(
    "📄 Upload Your CV Template",
    type=["docx", "jpg", "jpeg", "png", "pdf"],
    help="Upload your CV design (Word, Image, or PDF)"
)

# --- Optional photo upload ---
uploaded_photo = st.file_uploader("📸 Upload Your Photo (optional)", type=["jpg", "jpeg", "png"])

if uploaded_template:
    file_type = uploaded_template.name.split(".")[-1].lower()

    if file_type == "docx":
        # --- Word Template Mode ---
        doc = Document(uploaded_template)

        # Extract all text from document
        all_text = []
        for p in doc.paragraphs:
            all_text.append(p.text)
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    all_text.append(c.text)

        joined_text = "\n".join(all_text)
        placeholders = sorted(list(set(re.findall(r"{{(.*?)}}", joined_text))))

        if not placeholders:
            st.warning("⚠️ No placeholders found in this template.")
        else:
            st.success(f"✅ Found {len(placeholders)} placeholders.")
            st.caption(", ".join(placeholders))

            st.subheader("✏️ Enter Your Details")
            user_inputs = {}

            for field in placeholders:
                if field.lower() != "photo":
                    user_inputs[field] = st.text_area(
                        field.replace("_", " ").title(),
                        placeholder=f"Enter your {field.replace('_', ' ')} (leave blank to remove)",
                    )

            # --- Generate CV ---
            if st.button("🚀 Generate CV"):
                for p in doc.paragraphs:
                    for key, value in user_inputs.items():
                        tag = f"{{{{{key}}}}}"
                        if tag in p.text:
                            if value.strip():
                                p.text = p.text.replace(tag, value)
                            else:
                                p.text = ""

                    if "{{photo}}" in p.text.lower():
                        p.text = ""
                        if uploaded_photo:
                            run = p.add_run()
                            run.add_picture(BytesIO(uploaded_photo.read()), width=Inches(1.3))

                for t in doc.tables:
                    for r in t.rows:
                        for c in r.cells:
                            for key, value in user_inputs.items():
                                tag = f"{{{{{key}}}}}"
                                if tag in c.text:
                                    if value.strip():
                                        c.text = c.text.replace(tag, value)
                                    else:
                                        c.text = ""
                            if "{{photo}}" in c.text.lower():
                                c.text = ""
                                if uploaded_photo:
                                    paragraph = c.paragraphs[0]
                                    run = paragraph.add_run()
                                    run.add_picture(BytesIO(uploaded_photo.read()), width=Inches(1.3))

                for p in doc.paragraphs:
                    if not p.text.strip():
                        p.text = ""

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                st.success("🎉 CV generated successfully!")
                st.download_button(
                    label="📥 Download CV (Word)",
                    data=buffer,
                    file_name="Generated_CV.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

    elif file_type in ["jpg", "jpeg", "png", "pdf"]:
        # --- Image/PDF Template Mode ---
        st.image(uploaded_template, caption="Template Preview", use_container_width=True)
        st.info("📌 Note: This template will not auto-fill text. You can overlay details or attach this design in future versions.")
        if st.button("✅ Confirm Template Upload"):
            st.success("Template image uploaded successfully!")

    else:
        st.error("Unsupported file type. Please upload DOCX, JPG, JPEG, PNG, or PDF.")    # --- Detect placeholders ---
    joined_text = "\n".join(all_text)
    placeholders = sorted(list(set(re.findall(r"{{(.*?)}}", joined_text))))

    if not placeholders:
        st.warning("⚠️ No placeholders found in this template.")
    else:
        st.success(f"✅ Found {len(placeholders)} placeholders.")
        st.caption(", ".join(placeholders))

        st.subheader("✏️ Enter Your Details")
        user_inputs = {}

        for field in placeholders:
            if field.lower() != "photo":
                user_inputs[field] = st.text_area(
                    field.replace("_", " ").title(),
                    placeholder=f"Enter your {field.replace('_', ' ')} (leave blank to remove)",
                )

        # --- Generate CV ---
        if st.button("🚀 Generate CV"):
            for p in doc.paragraphs:
                for key, value in user_inputs.items():
                    tag = f"{{{{{key}}}}}"
                    if tag in p.text:
                        if value.strip():
                            p.text = p.text.replace(tag, value)
                        else:
                            p.text = ""

                if "{{photo}}" in p.text.lower():
                    p.text = ""
                    if uploaded_photo:
                        run = p.add_run()
                        run.add_picture(BytesIO(uploaded_photo.read()), width=Inches(1.3))

            # Handle placeholders in tables too
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        for key, value in user_inputs.items():
                            tag = f"{{{{{key}}}}}"
                            if tag in c.text:
                                if value.strip():
                                    c.text = c.text.replace(tag, value)
                                else:
                                    c.text = ""
                        if "{{photo}}" in c.text.lower():
                            c.text = ""
                            if uploaded_photo:
                                paragraph = c.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(uploaded_photo.read()), width=Inches(1.3))

            # Remove empty paragraphs
            for p in doc.paragraphs:
                if not p.text.strip():
                    p.text = ""

            # Save to memory buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.success("🎉 Your CV has been generated successfully!")
            st.download_button(
                label="📥 Download CV (Word)",
                data=buffer,
                file_name="Generated_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    if not placeholders:
        st.warning("⚠️ No placeholders found in this template.")
    else:
        st.success(f"✅ Found {len(placeholders)} placeholders.")
        st.caption(", ".join(placeholders))

        st.subheader("✏️ Enter Your Details")
        user_inputs = {}

        # --- Create text areas for each placeholder except 'photo'
        for field in placeholders:
            if field.lower() != "photo":
                user_inputs[field] = st.text_area(
                    field.replace("_", " ").title(),
                    placeholder=f"Enter your {field.replace('_', ' ')} (leave blank to remove)",
                )

        # --- Image upload only if {{photo}} exists ---
        photo_data = None
        if any(f.lower() == "photo" for f in placeholders):
            st.subheader("📸 Upload Profile Image")
            photo_file = st.file_uploader("Choose an image (JPG/PNG)", type=["jpg", "jpeg", "png"])
            if photo_file:
                # Read binary data
                photo_data = photo_file.read()

        # --- Generate Button ---
        if st.button("🚀 Generate CV"):
            for p in doc.paragraphs:
                for key, value in user_inputs.items():
                    tag = f"{{{{{key}}}}}"
                    if tag in p.text:
                        if value.strip():
                            p.text = p.text.replace(tag, value)
                        else:
                            p.text = ""

                # Handle photo placeholder in paragraph
                if "{{photo}}" in p.text.lower():
                    p.text = ""
                    if photo_data:
                        run = p.add_run()
                        run.add_picture(BytesIO(photo_data), width=Inches(1.3))

            # Handle photo + text placeholders inside tables
            for t in doc.tables:
                for r in t.rows:
                    for c in r.cells:
                        # Replace text placeholders
                        for key, value in user_inputs.items():
                            tag = f"{{{{{key}}}}}"
                            if tag in c.text:
                                if value.strip():
                                    c.text = c.text.replace(tag, value)
                                else:
                                    c.text = ""
                        # Insert photo inside tables
                        if "{{photo}}" in c.text.lower():
                            c.text = ""
                            if photo_data:
                                paragraph = c.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(photo_data), width=Inches(1.3))

            # Clean up empty paragraphs
            for p in doc.paragraphs:
                if not p.text.strip():
                    p.text = ""

            # --- Save and provide download ---
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.success("🎉 Your CV has been successfully generated!")
            st.download_button(
                label="📥 Download CV (Word)",
                data=buffer,
                file_name="Generated_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    if not placeholders:
        st.warning("⚠️ No placeholders found in the template.")
    else:
        st.success(f"✅ Found {len(placeholders)} placeholders.")
        st.caption(", ".join(placeholders))

        st.subheader("✏️ Enter Your Information")
        user_inputs = {}

        # --- Step 4: Input fields for text placeholders ---
        for field in placeholders:
            if field.lower() != "photo":  # text inputs for all except photo
                label = field.replace("_", " ").title()
                user_inputs[field] = st.text_area(label, placeholder=f"Enter your {label} (leave blank to remove section)")

        # --- Step 5: Upload Image if {{photo}} exists ---
        photo_data = None
        if "photo" in [f.lower() for f in placeholders]:
            st.subheader("📸 Upload Profile Photo")
            uploaded_image = st.file_uploader("Choose an image (JPG or PNG)", type=["jpg", "jpeg", "png"])
            if uploaded_image:
                photo_data = uploaded_image.read()

        # --- Step 6: Generate Button ---
        if st.button("🚀 Generate CV"):
            for p in doc.paragraphs:
                for key, value in user_inputs.items():
                    tag = f"{{{{{key}}}}}"
                    if tag in p.text:
                        if value.strip():
                            p.text = p.text.replace(tag, value)
                        else:
                            p.text = ""

                # Handle photo placeholder in paragraphs
                if "{{photo}}" in p.text.lower():
                    p.text = ""
                    if photo_data:
                        run = p.add_run()
                        run.add_picture(BytesIO(photo_data), width=Inches(1.2))

            # Handle photo placeholders inside tables too
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        # Replace text placeholders
                        for key, value in user_inputs.items():
                            tag = f"{{{{{key}}}}}"
                            if tag in cell.text:
                                if value.strip():
                                    cell.text = cell.text.replace(tag, value)
                                else:
                                    cell.text = ""
                        # Handle {{photo}} in tables
                        if "{{photo}}" in cell.text.lower():
                            cell.text = ""
                            if photo_data:
                                paragraph = cell.paragraphs[0]
                                run = paragraph.add_run()
                                run.add_picture(BytesIO(photo_data), width=Inches(1.2))

            # Remove empty paragraphs
            for p in doc.paragraphs:
                if not p.text.strip():
                    p.text = ""

            # --- Step 7: Download File ---
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            st.success("🎉 Your CV is ready for download!")
            st.download_button(
                label="📥 Download CV (Word)",
                data=buffer,
                file_name="Generated_CV.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
